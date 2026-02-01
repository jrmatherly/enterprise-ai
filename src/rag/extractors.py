"""Document text extraction for various file types.

Supports: PDF, DOCX, TXT, Markdown
"""

from abc import ABC, abstractmethod
from io import BytesIO
from typing import Optional
import re


class ExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


class TextExtractor(ABC):
    """Base class for text extractors."""
    
    @abstractmethod
    def extract(self, content: bytes) -> str:
        """Extract text from document content."""
        pass
    
    @abstractmethod
    def supported_types(self) -> list[str]:
        """Return list of supported MIME types."""
        pass


class PlainTextExtractor(TextExtractor):
    """Extract text from plain text and markdown files."""
    
    def extract(self, content: bytes) -> str:
        """Decode bytes to text."""
        # Try common encodings
        for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort: ignore errors
        return content.decode("utf-8", errors="ignore")
    
    def supported_types(self) -> list[str]:
        return ["text/plain", "text/markdown", "text/x-markdown"]


class PDFExtractor(TextExtractor):
    """Extract text from PDF files using pypdf."""
    
    def extract(self, content: bytes) -> str:
        """Extract text from all pages of a PDF."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ExtractionError("pypdf not installed. Run: pip install pypdf")
        
        try:
            reader = PdfReader(BytesIO(content))
            
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                except Exception as e:
                    text_parts.append(f"[Page {page_num}] (extraction failed: {e})")
            
            if not text_parts:
                raise ExtractionError("No text could be extracted from PDF")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            if "extraction failed" in str(e):
                raise
            raise ExtractionError(f"PDF extraction failed: {e}")
    
    def supported_types(self) -> list[str]:
        return ["application/pdf"]


class DOCXExtractor(TextExtractor):
    """Extract text from Word documents using python-docx."""
    
    def extract(self, content: bytes) -> str:
        """Extract text from DOCX, preserving paragraph structure."""
        try:
            from docx import Document
        except ImportError:
            raise ExtractionError("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = Document(BytesIO(content))
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Preserve heading structure
                    if para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading ", "")
                        try:
                            level_num = int(level)
                            prefix = "#" * level_num + " "
                        except ValueError:
                            prefix = "# "
                        text_parts.append(f"{prefix}{text}")
                    else:
                        text_parts.append(text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables, 1):
                table_text = [f"\n[Table {table_idx}]"]
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.replace("|", "").strip():
                        table_text.append(row_text)
                if len(table_text) > 1:
                    text_parts.extend(table_text)
            
            if not text_parts:
                raise ExtractionError("No text could be extracted from DOCX")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            if "extraction failed" in str(e) or "No text" in str(e):
                raise
            raise ExtractionError(f"DOCX extraction failed: {e}")
    
    def supported_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]


class DocumentExtractor:
    """Unified document extractor that delegates to specific extractors."""
    
    def __init__(self):
        self.extractors: list[TextExtractor] = [
            PlainTextExtractor(),
            PDFExtractor(),
            DOCXExtractor(),
        ]
        
        # Build MIME type mapping
        self._mime_map: dict[str, TextExtractor] = {}
        for extractor in self.extractors:
            for mime_type in extractor.supported_types():
                self._mime_map[mime_type] = extractor
    
    def supports(self, mime_type: str) -> bool:
        """Check if a MIME type is supported."""
        return mime_type in self._mime_map
    
    def supported_types(self) -> list[str]:
        """Get all supported MIME types."""
        return list(self._mime_map.keys())
    
    def extract(self, content: bytes, mime_type: str) -> str:
        """Extract text from document based on MIME type.
        
        Args:
            content: Raw document bytes
            mime_type: Document MIME type
            
        Returns:
            Extracted text
            
        Raises:
            ExtractionError: If extraction fails or type not supported
        """
        extractor = self._mime_map.get(mime_type)
        
        if not extractor:
            raise ExtractionError(
                f"Unsupported file type: {mime_type}. "
                f"Supported types: {', '.join(self.supported_types())}"
            )
        
        text = extractor.extract(content)
        
        # Clean up extracted text
        text = self._clean_text(text)
        
        return text
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        
        # Remove excessive whitespace
        text = re.sub(r"[ \t]+", " ", text)
        
        # Remove excessive newlines (more than 2)
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        # Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)
        
        return text.strip()


# Singleton instance
_extractor: Optional[DocumentExtractor] = None


def get_extractor() -> DocumentExtractor:
    """Get or create the global DocumentExtractor instance."""
    global _extractor
    if _extractor is None:
        _extractor = DocumentExtractor()
    return _extractor
